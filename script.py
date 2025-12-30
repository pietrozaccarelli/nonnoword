import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import json
import os
import uuid
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ToolTip:
    """Class to create a hover tooltip for a widget."""
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip_window = None
        self.widget.bind("<Enter>", self.show_tip)
        self.widget.bind("<Leave>", self.hide_tip)

    def show_tip(self, event=None):
        if self.tip_window or not self.text:
            return
        x, y, _cx, cy = self.widget.bbox("insert")
        x = x + self.widget.winfo_rootx() + 25
        y = y + cy + self.widget.winfo_rooty() + 25
        self.tip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(1)
        tw.wm_geometry(f"+{x}+{y}")
        label = tk.Label(tw, text=self.text, justify=tk.LEFT,
                         background="#ffffe0", relief=tk.SOLID, borderwidth=1,
                         font=("tahoma", "8", "normal"))
        label.pack(ipadx=1)

    def hide_tip(self, event=None):
        tw = self.tip_window
        self.tip_window = None
        if tw:
            tw.destroy()

class WordEmulator:
    def __init__(self, root):
        self.root = root
        self.root.title("Python Word Emulator")
        self.root.geometry("1300x850")

        # --- State and Paths ---
        self.base_dir = os.path.dirname(os.path.abspath(__file__))
        self.index_dir = os.path.join(self.base_dir, "index")
        self.index_path = os.path.join(self.index_dir, "index.json")
        self.backups_dir = os.path.join(self.base_dir, "backups")
        
        self.periodic_backup_active = tk.BooleanVar(value=True)
        self.backup_duration_var = tk.StringVar(value="2") 
        self.countdown_seconds = 120 
        self.font_size_var = tk.IntVar(value=12)
        
        self.current_style = {"bold": False, "italic": False, "underline": False, "size": 12}
        self.current_file_path = None
        self.doc_id = str(uuid.uuid4().hex)[:12] 
        self.file_name = self.doc_id 
        
        # Fingerprints
        self.last_saved_fingerprint = None  # Compares against main file
        self.last_backup_fingerprint = None # Compares against last backup iteration

        self.ensure_dirs()
        self.setup_ui()
        self.update_window_title()
        self.start_timer_loop()

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def ensure_dirs(self):
        os.makedirs(self.index_dir, exist_ok=True)
        os.makedirs(self.backups_dir, exist_ok=True)
        if not os.path.exists(self.index_path):
            with open(self.index_path, "w") as f: json.dump({}, f)

    def register_in_index(self, doc_id, name):
        with open(self.index_path, "r") as f: 
            try: index_data = json.load(f)
            except: index_data = {}
        index_data[doc_id] = name
        with open(self.index_path, "w") as f: json.dump(index_data, f, indent=4)

    def get_name_from_index(self, doc_id):
        with open(self.index_path, "r") as f:
            try: return json.load(f).get(doc_id)
            except: return None

    def setup_ui(self):
        self.toolbar = ttk.Frame(self.root); self.toolbar.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)

        # File Group
        file_g = self.create_tool_group("File", tk.LEFT)
        btn_new = ttk.Button(file_g, text="\U0001f4c4 New", command=self.new_file, takefocus=False)
        btn_new.pack(side=tk.LEFT, padx=2)
        ToolTip(btn_new, "Save current and start a blank document (Ctrl+N)")

        btn_load = ttk.Button(file_g, text="\U0001f4c2 Load", command=self.load_file, takefocus=False)
        btn_load.pack(side=tk.LEFT, padx=2)
        ToolTip(btn_load, "Open an existing document")

        btn_save = ttk.Button(file_g, text="\U0001f4be Save", command=self.save_file, takefocus=False)
        btn_save.pack(side=tk.LEFT, padx=2)
        ToolTip(btn_save, "Overwrite current file (Ctrl+S)")

        btn_save_as = ttk.Button(file_g, text="\U0001f4dd Save As", command=self.save_as_file, takefocus=False)
        btn_save_as.pack(side=tk.LEFT, padx=2)
        ToolTip(btn_save_as, "Save as a new file identity")

        # Size Group
        size_g = self.create_tool_group("Size", tk.LEFT)
        self.size_box = ttk.Combobox(size_g, textvariable=self.font_size_var, values=list(range(8, 73, 2)), width=3)
        self.size_box.pack(side=tk.LEFT, padx=2)
        self.size_box.bind("<<ComboboxSelected>>", lambda e: self.apply_formatting("size"))
        ttk.Button(size_g, text="+", width=2, command=lambda: self.adjust_font_size(1), takefocus=False).pack(side=tk.LEFT)
        ttk.Button(size_g, text="-", width=2, command=lambda: self.adjust_font_size(-1), takefocus=False).pack(side=tk.LEFT)

        # Format Group
        format_g = self.create_tool_group("Format", tk.LEFT)
        btn_p = {"width": 3, "relief": "raised", "bd": 2}
        self.btn_bold = tk.Button(format_g, text="B", font=("Arial", 9, "bold"), command=lambda: self.apply_formatting("bold"), **btn_p)
        self.btn_bold.pack(side=tk.LEFT, padx=1)
        self.btn_italic = tk.Button(format_g, text="I", font=("Arial", 9, "italic"), command=lambda: self.apply_formatting("italic"), **btn_p)
        self.btn_italic.pack(side=tk.LEFT, padx=1)
        self.btn_underline = tk.Button(format_g, text="U", font=("Arial", 9, "underline"), command=lambda: self.apply_formatting("underline"), **btn_p)
        self.btn_underline.pack(side=tk.LEFT, padx=1)

        # Alignment
        align_g = self.create_tool_group("Alignment", tk.LEFT)
        ttk.Button(align_g, text="Left", width=5, command=lambda: self.set_alignment("left"), takefocus=False).pack(side=tk.LEFT)
        ttk.Button(align_g, text="Center", width=7, command=lambda: self.set_alignment("center"), takefocus=False).pack(side=tk.LEFT)
        ttk.Button(align_g, text="Right", width=6, command=lambda: self.set_alignment("right"), takefocus=False).pack(side=tk.LEFT)

        # Backup Group
        backup_g = self.create_tool_group("Backup Settings", tk.RIGHT)
        ttk.Checkbutton(backup_g, text="Periodic", variable=self.periodic_backup_active, takefocus=False).pack(side=tk.LEFT, padx=5)
        ttk.Entry(backup_g, textvariable=self.backup_duration_var, width=3).pack(side=tk.LEFT)
        ttk.Label(backup_g, text="min").pack(side=tk.LEFT, padx=2)
        btn_set = ttk.Button(backup_g, text="Set", width=4, command=self.reset_countdown, takefocus=False)
        btn_set.pack(side=tk.LEFT, padx=2)
        self.lbl_countdown = ttk.Label(backup_g, text="Next: 02:00", font=("Consolas", 10, "bold"), foreground="#d9534f")
        self.lbl_countdown.pack(side=tk.LEFT, padx=10)

        # Editor
        ed_fr = ttk.Frame(self.root); ed_fr.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        sb = ttk.Scrollbar(ed_fr); sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.text_area = tk.Text(ed_fr, undo=True, font=("Calibri", 12), wrap="word", yscrollcommand=sb.set)
        self.text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.config(command=self.text_area.yview)
        
        self.text_area.bind("<<Selection>>", self.detect_format_at_cursor)
        self.text_area.bind("<Key>", self.on_key_press)
        self.text_area.bind("<ButtonRelease-1>", self.detect_format_at_cursor)
        self.text_area.bind("<KeyRelease>", self.detect_format_at_cursor)
        self.text_area.tag_configure("left", justify="left")
        self.text_area.tag_configure("center", justify="center")
        self.text_area.tag_configure("right", justify="right")

        self.setup_bindings()
        self.status = ttk.Label(self.root, text="Ready", relief=tk.SUNKEN, anchor=tk.W); self.status.pack(side=tk.BOTTOM, fill=tk.X)

    def create_tool_group(self, text, side):
        c = ttk.Frame(self.toolbar); c.pack(side=side, padx=10)
        ttk.Label(c, text=text, font=("Segoe UI", 9, "bold")).pack(side=tk.TOP, anchor=tk.W)
        r = ttk.Frame(c); r.pack(side=tk.TOP); return r

    def setup_bindings(self):
        self.text_area.bind("<Control-n>", lambda e: self.new_file() or "break")
        self.text_area.bind("<Control-N>", lambda e: self.new_file() or "break")
        self.text_area.bind("<Control-b>", lambda e: self.apply_formatting("bold") or "break")
        self.text_area.bind("<Control-i>", lambda e: self.apply_formatting("italic") or "break")
        self.text_area.bind("<Control-u>", lambda e: self.apply_formatting("underline") or "break")
        self.text_area.bind("<Control-s>", lambda e: self.save_file() or "break")

    def update_window_title(self):
        name = self.file_name if self.file_name else "New Document"
        self.root.title(f"Word Emulator - {name}")

    def get_fingerprint(self):
        raw_dump = self.text_area.dump("1.0", "end-1c", text=True, tag=True)
        return [item for item in raw_dump if item[0] in ('text', 'tagon', 'tagoff')]

    def start_timer_loop(self):
        if self.periodic_backup_active.get():
            self.countdown_seconds -= 1
            if self.countdown_seconds <= 0:
                self.perform_backup()
                self.reset_countdown()
        m, s = divmod(max(0, self.countdown_seconds), 60)
        self.lbl_countdown.config(text=f"Next: {m:02d}:{s:02d}")
        self.root.after(1000, self.start_timer_loop)

    def reset_countdown(self):
        try:
            val = int(self.backup_duration_var.get())
            self.countdown_seconds = val * 60
        except:
            self.backup_duration_var.set("2"); self.countdown_seconds = 120

    def perform_backup(self):
        current_data = self.get_fingerprint()
        if not self.text_area.get("1.0", "end-1c").strip(): return
        
        # SMART BACKUP: Only skip if identical to the LAST BACKUP file
        if current_data == self.last_backup_fingerprint: return

        target = os.path.join(self.backups_dir, self.file_name)
        os.makedirs(target, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d%H%M")
        self.write_docx(os.path.join(target, f"{self.file_name}_{ts}.docx"))
        
        self.last_backup_fingerprint = current_data
        self.status.config(text=f"Auto-backup saved: {ts[-4:-2]}:{ts[-2:]}")

    def sync_structure(self, new_full_path):
        """Renames backup folder and all internal backup files to match new filename."""
        new_name_base = os.path.splitext(os.path.basename(new_full_path))[0]
        old_name_base = self.file_name
        
        if old_name_base == new_name_base: return

        old_dir = os.path.join(self.backups_dir, old_name_base)
        new_dir = os.path.join(self.backups_dir, new_name_base)

        if os.path.exists(old_dir):
            for f in os.listdir(old_dir):
                # Regex matches [Prefix]_[12 Digits].docx
                match = re.match(r"(.+)_(\d{12})\.docx$", f)
                if match:
                    new_f = f"{new_name_base}_{match.group(2)}.docx"
                    try: os.rename(os.path.join(old_dir, f), os.path.join(old_dir, new_f))
                    except: pass
            try:
                if not os.path.exists(new_dir): os.rename(old_dir, new_dir)
            except: pass

        self.current_file_path = new_full_path
        self.file_name = new_name_base
        self.update_window_title()

    def new_file(self):
        if self.text_area.get("1.0", "end-1c").strip(): self.save_file()
        self.doc_id = str(uuid.uuid4().hex)[:12]
        self.file_name = self.doc_id
        self.current_file_path = None
        self.last_saved_fingerprint = None
        self.last_backup_fingerprint = None
        self.text_area.delete("1.0", tk.END)
        self.current_style = {"bold": False, "italic": False, "underline": False, "size": 12}
        self.font_size_var.set(12)
        self.update_ui_buttons(); self.update_window_title()
        self.status.config(text="New document created")
        self.text_area.focus_set()

    def load_file(self):
        path = filedialog.askopenfilename(filetypes=[("Word", "*.docx")])
        if not path: return
        doc = Document(path)
        loaded_id = doc.core_properties.identifier
        new_disk_name = os.path.splitext(os.path.basename(path))[0]
        
        if loaded_id:
            old_name = self.get_name_from_index(loaded_id)
            if old_name: self.file_name = old_name
            self.doc_id = loaded_id
        else:
            self.doc_id = str(uuid.uuid4().hex)[:12]
            self.file_name = new_disk_name

        self.sync_structure(path) # Detects if filename changed on disk
        self.register_in_index(self.doc_id, self.file_name)

        t_meta = doc.core_properties.comments
        self.backup_duration_var.set(t_meta if (t_meta and t_meta.isdigit()) else "2")
        self.reset_countdown()
        self.text_area.delete("1.0", tk.END)
        for p in doc.paragraphs:
            p_st = self.text_area.index("end-1c")
            for run in p.runs:
                if not run.text: continue
                r_st = self.text_area.index("end-1c"); self.text_area.insert(tk.END, run.text); r_en = self.text_area.index("end-1c")
                sz = int(run.font.size.pt) if run.font.size else 12
                self.current_style.update({"bold": run.bold, "italic": run.italic, "underline": run.underline, "size": sz})
                self.apply_style_to_range(r_st, r_en)
            al = "left"
            if p.alignment == WD_ALIGN_PARAGRAPH.CENTER: al = "center"
            elif p.alignment == WD_ALIGN_PARAGRAPH.RIGHT: al = "right"
            self.text_area.tag_add(al, p_st, "end-1c"); self.text_area.insert(tk.END, "\n")
        
        self.last_saved_fingerprint = self.get_fingerprint()
        self.last_backup_fingerprint = None 
        self.text_area.focus_set()

    def write_docx(self, path):
        doc = Document(); doc.core_properties.identifier = self.doc_id 
        doc.core_properties.comments = self.backup_duration_var.get() 
        content = self.text_area.dump("1.0", "end-1c", text=True, tag=True)
        active_tags, p = set(), doc.add_paragraph()
        a_map = {"center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT, "left": WD_ALIGN_PARAGRAPH.LEFT}
        for type, val, idx in content:
            if type == "tagon": active_tags.add(val)
            elif type == "tagoff": active_tags.discard(val)
            elif type == "text":
                if val == "\n": p = doc.add_paragraph(); continue
                run = p.add_run(val)
                run.bold, run.italic, run.underline = "bold" in active_tags, "italic" in active_tags, "underline" in active_tags
                sz = 12
                for t in active_tags:
                    if t.startswith("sz_"): sz = int(t.split("_")[1])
                run.font.size = Pt(sz)
                for a, wd_a in a_map.items():
                    if a in active_tags: p.alignment = wd_a
        doc.save(path)

    def save_file(self):
        if self.current_file_path:
            self.write_docx(self.current_file_path); self.register_in_index(self.doc_id, self.file_name)
            self.last_saved_fingerprint = self.get_fingerprint(); self.status.config(text="Updated")
            return True
        return self.save_as_file()

    def save_as_file(self):
        p = filedialog.asksaveasfilename(initialfile=f"{self.file_name}.docx", defaultextension=".docx", filetypes=[("Word", "*.docx")])
        if not p: return False
        self.sync_structure(p) 
        self.doc_id = str(uuid.uuid4().hex)[:12] 
        self.write_docx(self.current_file_path); self.register_in_index(self.doc_id, self.file_name)
        self.last_saved_fingerprint = self.get_fingerprint(); self.last_backup_fingerprint = None
        return True

    # --- Formatting Helpers ---
    def update_ui_buttons(self):
        self.btn_bold.config(bg="#add8e6" if self.current_style["bold"] else "SystemButtonFace", relief="sunken" if self.current_style["bold"] else "raised")
        self.btn_italic.config(bg="#add8e6" if self.current_style["italic"] else "SystemButtonFace", relief="sunken" if self.current_style["italic"] else "raised")
        self.btn_underline.config(bg="#add8e6" if self.current_style["underline"] else "SystemButtonFace", relief="sunken" if self.current_style["underline"] else "raised")

    def detect_format_at_cursor(self, event=None):
        try: idx = self.text_area.index("sel.first")
        except tk.TclError:
            idx = self.text_area.index("insert - 1c")
            if self.text_area.compare("insert", "==", "1.0"): return
        tags = self.text_area.tag_names(idx)
        self.current_style.update({"bold": "bold" in tags, "italic": "italic" in tags, "underline": "underline" in tags})
        for t in tags:
            if t.startswith("sz_"):
                self.current_style["size"] = int(t.split("_")[1]); self.font_size_var.set(self.current_style["size"]); break
        self.update_ui_buttons()

    def apply_formatting(self, style_type):
        try:
            st, en = self.text_area.index("sel.first"), self.text_area.index("sel.last")
            self.apply_style_to_range(st, en, toggle_type=style_type)
        except tk.TclError:
            if style_type == "bold": self.current_style["bold"] = not self.current_style["bold"]
            elif style_type == "italic": self.current_style["italic"] = not self.current_style["italic"]
            elif style_type == "underline": self.current_style["underline"] = not self.current_style["underline"]
            elif style_type == "size": self.current_style["size"] = self.font_size_var.get()
        self.detect_format_at_cursor(); self.text_area.focus_set()

    def apply_style_to_range(self, start, end, toggle_type=None):
        curr = self.text_area.index(start)
        while self.text_area.compare(curr, "<", end):
            tags = self.text_area.tag_names(curr)
            is_b = not ("bold" in tags) if toggle_type == "bold" else (self.current_style["bold"] if not toggle_type else "bold" in tags)
            is_i = not ("italic" in tags) if toggle_type == "italic" else (self.current_style["italic"] if not toggle_type else "italic" in tags)
            is_u = not ("underline" in tags) if toggle_type == "underline" else (self.current_style["underline"] if not toggle_type else "underline" in tags)
            sz = self.font_size_var.get() if toggle_type == "size" else (self.current_style["size"] if not toggle_type else self.get_size_at(curr))
            
            for t in tags:
                if t in ["bold", "italic", "underline"] or t.startswith("sz_") or t.startswith("comp_"):
                    self.text_area.tag_remove(t, curr)
            if is_b: self.text_area.tag_add("bold", curr)
            if is_i: self.text_area.tag_add("italic", curr)
            if is_u: self.text_area.tag_add("underline", curr)
            self.text_area.tag_add(f"sz_{sz}", curr)
            comp = f"comp_{sz}_{'b' if is_b else ''}{'i' if is_i else ''}"
            st_l = ["bold" if is_b else None, "italic" if is_i else None]
            self.text_area.tag_configure(comp, font=("Calibri", sz, " ".join(filter(None, st_l))))
            self.text_area.tag_configure("underline", underline=True)
            self.text_area.tag_add(comp, curr); curr = self.text_area.index(f"{curr} + 1c")

    def get_size_at(self, index):
        for t in self.text_area.tag_names(index):
            if t.startswith("sz_"): return int(t.split("_")[1])
        return 12

    def on_key_press(self, event):
        if len(event.char) > 0 and event.char.isprintable():
            self.root.after(1, lambda: self.apply_style_to_range("insert - 1c", "insert"))

    def set_alignment(self, align):
        try:
            try: st, en = self.text_area.index("sel.first"), self.text_area.index("sel.last")
            except: st, en = self.text_area.index("insert linestart"), self.text_area.index("insert lineend")
            for a in ["left", "center", "right"]: self.text_area.tag_remove(a, st, en)
            self.text_area.tag_add(align, st, en)
        except: pass
        self.text_area.focus_set()

    def on_closing(self):
        if self.text_area.get("1.0", "end-1c").strip():
            if self.save_file(): self.root.destroy()
        else: self.root.destroy()

if __name__ == "__main__":
    root = tk.Tk(); app = WordEmulator(root); root.mainloop()